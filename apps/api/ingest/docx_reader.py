"""
DOCX Reader Module

Extracts paragraphs from Arabic .docx files with stable anchors.
Uses python-docx for parsing and extracts paragraph IDs when available.
"""

import hashlib
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn
from lxml import etree


@dataclass
class ParsedParagraph:
    """
    A parsed paragraph with stable anchor information.

    Attributes:
        para_index: Zero-based paragraph index in document order.
        text: The paragraph text content.
        style_name: Word style name (e.g., "Heading 1").
        runs: List of run texts (for detecting bold/italic markers).
        para_id: Word paragraph ID (w14:paraId) if available.
        text_hash: SHA1 hash of normalized text for stable fallback anchor.
    """

    para_index: int
    text: str
    style_name: Optional[str] = None
    runs: list[str] = field(default_factory=list)
    para_id: Optional[str] = None
    text_hash: str = ""

    def __post_init__(self):
        """Compute text hash if not provided."""
        if not self.text_hash and self.text:
            normalized = self._normalize_for_hash(self.text)
            self.text_hash = hashlib.sha1(normalized.encode("utf-8")).hexdigest()[:12]

    @staticmethod
    def _normalize_for_hash(text: str) -> str:
        """Normalize text for hash computation."""
        # Remove diacritics (tashkeel)
        text = re.sub(r"[\u064B-\u065F\u0670]", "", text)
        # Normalize whitespace
        text = " ".join(text.split())
        return text.strip()

    @property
    def anchor_id(self) -> str:
        """
        Get the stable anchor ID.

        Prefers w14:paraId if available, falls back to para_index + text_hash.
        """
        if self.para_id:
            return f"pid_{self.para_id}"
        return f"p{self.para_index}_{self.text_hash}"


@dataclass
class ParsedDocument:
    """
    A fully parsed document with metadata.

    Attributes:
        file_name: Original file name.
        file_hash: SHA256 hash of the file content.
        paragraphs: List of parsed paragraphs.
        total_paragraphs: Total number of paragraphs.
    """

    file_name: str
    file_hash: str
    paragraphs: list[ParsedParagraph]
    total_paragraphs: int = 0

    def __post_init__(self):
        """Set total paragraphs if not provided."""
        if not self.total_paragraphs:
            self.total_paragraphs = len(self.paragraphs)


class DocxReader:
    """
    Reader for Arabic .docx files with stable anchor extraction.

    This reader:
    1. Parses all paragraphs in document order
    2. Extracts Word paragraph IDs (w14:paraId) when available
    3. Computes fallback anchors using para_index + text_hash
    4. Preserves style information and run-level formatting
    """

    # Namespace for Word 2010+ paragraph IDs
    W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

    def __init__(self):
        """Initialize the DOCX reader."""
        self._para_ids: dict[int, str] = {}

    def read(self, file_path: str | Path) -> ParsedDocument:
        """
        Read and parse a .docx file.

        Args:
            file_path: Path to the .docx file.

        Returns:
            ParsedDocument: The parsed document with all paragraphs.

        Raises:
            FileNotFoundError: If the file does not exist.
            ValueError: If the file is not a valid .docx file.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if not path.suffix.lower() == ".docx":
            raise ValueError(f"Not a .docx file: {file_path}")

        # Compute file hash
        file_hash = self._compute_file_hash(path)

        # Extract paragraph IDs from XML first
        self._extract_paragraph_ids(path)

        # Parse document with python-docx
        doc = Document(str(path))
        paragraphs = self._parse_paragraphs(doc)

        return ParsedDocument(
            file_name=path.name,
            file_hash=file_hash,
            paragraphs=paragraphs,
        )

    def read_bytes(self, content: bytes, file_name: str) -> ParsedDocument:
        """
        Read and parse .docx content from bytes.

        Args:
            content: The .docx file content as bytes.
            file_name: Name to use for the document.

        Returns:
            ParsedDocument: The parsed document.
        """
        import io

        # Compute hash from content
        file_hash = hashlib.sha256(content).hexdigest()

        # Create a BytesIO object for python-docx
        doc_stream = io.BytesIO(content)

        # Extract paragraph IDs from the stream
        self._extract_paragraph_ids_from_stream(content)

        # Parse document
        doc = Document(doc_stream)
        paragraphs = self._parse_paragraphs(doc)

        return ParsedDocument(
            file_name=file_name,
            file_hash=file_hash,
            paragraphs=paragraphs,
        )

    def _compute_file_hash(self, path: Path) -> str:
        """Compute SHA256 hash of file content."""
        sha256 = hashlib.sha256()
        with open(path, "rb") as f:
            for chunk in iter(lambda: f.read(8192), b""):
                sha256.update(chunk)
        return sha256.hexdigest()

    def _extract_paragraph_ids(self, path: Path) -> None:
        """
        Extract w14:paraId attributes from the document XML.

        Word 2010+ stores paragraph IDs in the document.xml as w14:paraId.
        """
        self._para_ids = {}

        try:
            with ZipFile(path, "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    return

                xml_content = zf.read("word/document.xml")
                self._parse_para_ids_from_xml(xml_content)
        except Exception:
            # If extraction fails, we'll use fallback anchors
            pass

    def _extract_paragraph_ids_from_stream(self, content: bytes) -> None:
        """Extract paragraph IDs from bytes content."""
        import io

        self._para_ids = {}

        try:
            with ZipFile(io.BytesIO(content), "r") as zf:
                if "word/document.xml" not in zf.namelist():
                    return

                xml_content = zf.read("word/document.xml")
                self._parse_para_ids_from_xml(xml_content)
        except Exception:
            pass

    def _parse_para_ids_from_xml(self, xml_content: bytes) -> None:
        """Parse paragraph IDs from document XML."""
        try:
            root = etree.fromstring(xml_content)

            # Find all w:p elements (paragraphs)
            nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
            paragraphs = root.findall(".//w:p", namespaces=nsmap)

            for idx, para in enumerate(paragraphs):
                # Look for w14:paraId attribute
                para_id = para.get(qn("w14:paraId"))
                if para_id:
                    self._para_ids[idx] = para_id
        except Exception:
            pass

    def _parse_paragraphs(self, doc: Document) -> list[ParsedParagraph]:
        """
        Parse all paragraphs from a python-docx Document.

        Args:
            doc: The python-docx Document object.

        Returns:
            List of ParsedParagraph objects.
        """
        paragraphs = []

        for idx, para in enumerate(doc.paragraphs):
            # Get text content
            text = para.text.strip()

            # Get style name
            style_name = None
            if para.style:
                style_name = para.style.name

            # Get run texts (for detecting formatting)
            runs = [run.text for run in para.runs if run.text]

            # Get paragraph ID if available
            para_id = self._para_ids.get(idx)

            parsed = ParsedParagraph(
                para_index=idx,
                text=text,
                style_name=style_name,
                runs=runs,
                para_id=para_id,
            )

            paragraphs.append(parsed)

        return paragraphs


def create_source_anchor(
    doc_hash: str,
    paragraph: ParsedParagraph,
    anchor_range: Optional[str] = None,
) -> dict:
    """
    Create a source anchor dictionary from a parsed paragraph.

    This follows the stable anchor scheme defined in the plan:
    - source_doc_id: Derived from file hash
    - anchor_type: docx_para
    - anchor_id: Prefer w14:paraId, fallback to para_index + text_hash
    - anchor_range: Optional, for multi-paragraph chunks

    Args:
        doc_hash: The document's SHA256 hash.
        paragraph: The parsed paragraph.
        anchor_range: Optional range for multi-paragraph chunks.

    Returns:
        Dictionary with anchor information.
    """
    return {
        "source_doc_id": f"DOC_{doc_hash[:16]}",
        "anchor_type": "docx_para",
        "anchor_id": paragraph.anchor_id,
        "anchor_range": anchor_range,
    }

